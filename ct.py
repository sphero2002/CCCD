from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Length, Pt
import json

def twips_to_pixels(twips):
    return int(twips * (96 / 1440))  # 96 DPI là tiêu chuẩn cho màn hình

class TableHelper:
    def __init__(self, table):
        self.table = table
        self.tbl_pr = table._element.tblPr  # Truy cập phần tblPr trong XML

    def get_alignment(self):
        jc = self.tbl_pr.find(qn('w:jc'))
        if jc is not None:
            val = jc.get(qn('w:val'))
            return val
        return None

    def set_alignment(self, alignment):
        jc = self.tbl_pr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            self.tbl_pr.append(jc)
        jc.set(qn('w:val'), alignment)

    def get_allow_autofit(self):
        tbl_layout = self.tbl_pr.find(qn('w:tblLayout'))
        if tbl_layout is not None:
            val = tbl_layout.get(qn('w:type'))
            return val == 'autofit'
        return True  # Mặc định là True

    def set_allow_autofit(self, allow):
        tbl_layout = self.tbl_pr.find(qn('w:tblLayout'))
        if tbl_layout is None:
            tbl_layout = OxmlElement('w:tblLayout')
            self.tbl_pr.append(tbl_layout)
        tbl_layout.set(qn('w:type'), 'autofit' if allow else 'fixed')

    def get_table_width(self):
        tbl_w = self.tbl_pr.find(qn('w:tblW'))
        if tbl_w is not None:
            return {
                'type': tbl_w.get(qn('w:type')),
                'width': twips_to_pixels(int(tbl_w.get(qn('w:w'), '0')))
            }
        return {'type': 'auto', 'width': 0}

    def get_table_borders(self):
        tbl_borders = self.tbl_pr.find(qn('w:tblBorders'))
        borders = {}
        if tbl_borders is not None:
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                element = tbl_borders.find(qn(f'w:{border}'))
                if element is not None:
                    borders[border] = {
                        'val': element.get(qn('w:val')),
                        'sz': element.get(qn('w:sz')),
                        'color': element.get(qn('w:color')),
                        'space': element.get(qn('w:space')),
                        'shadow': element.get(qn('w:shadow'))
                    }
        return borders

    def get_table_look(self):
        tbl_look = self.tbl_pr.find(qn('w:tblLook'))
        look = {}
        if tbl_look is not None:
            for attr in tbl_look.attrib:
                look[attr.replace('w:', '')] = tbl_look.get(attr)
        return look

    def get_columns_width(self):
        columns = []
        for col in self.table.columns:
            width = col.width
            if width:
                columns.append(twips_to_pixels(width.twips))
            else:
                columns.append(None)
        return columns

    def get_rows_height(self):
        rows = []
        for row in self.table.rows:
            height = row.height
            if height:
                height_pixels = twips_to_pixels(height.twips)
                height_rule = row.height_rule
                rows.append({'height': height_pixels, 'height_rule': height_rule})
            else:
                rows.append({'height': None, 'height_rule': None})
        return rows

    def get_cells_properties(self):
        cells_props = []
        for row in self.table.rows:
            row_props = []
            for cell in row.cells:
                cell_prop = {}
                tc = cell._element  # Đã là CT_Tc
                tc_pr = tc.find(qn('w:tcPr'))
                if tc_pr is not None:
                    # Chiều rộng ô
                    tc_w = tc_pr.find(qn('w:tcW'))
                    if tc_w is not None:
                        cell_prop['width'] = {
                            'type': tc_w.get(qn('w:type')),
                            'width': twips_to_pixels(int(tc_w.get(qn('w:w'), '0')))
                        }
                    # Căn chỉnh dọc
                    v_align = tc_pr.find(qn('w:vAlign'))
                    if v_align is not None:
                        cell_prop['vertical_alignment'] = v_align.get(qn('w:val'))
                    else:
                        cell_prop['vertical_alignment'] = None
                    # Biên ô
                    cell_borders = tc_pr.find(qn('w:tcBorders'))
                    borders = {}
                    if cell_borders is not None:
                        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            element = cell_borders.find(qn(f'w:{border}'))
                            if element is not None:
                                borders[border] = {
                                    'val': element.get(qn('w:val')),
                                    'sz': element.get(qn('w:sz')),
                                    'color': element.get(qn('w:color')),
                                    'space': element.get(qn('w:space')),
                                    'shadow': element.get(qn('w:shadow'))
                                }
                    cell_prop['borders'] = borders
                    # Hợp nhất ô
                    grid_span = tc_pr.find(qn('w:gridSpan'))
                    if grid_span is not None:
                        cell_prop['grid_span'] = int(grid_span.get(qn('w:val'), '1'))
                    else:
                        cell_prop['grid_span'] = 1
                    v_merge = tc_pr.find(qn('w:vMerge'))
                    if v_merge is not None:
                        v_merge_val = v_merge.get(qn('w:val'))
                        cell_prop['v_merge'] = v_merge_val if v_merge_val else 'continue'
                    else:
                        cell_prop['v_merge'] = None
                    # Bạn có thể thêm xử lý các thuộc tính khác như shading, noWrap, etc. ở đây
                row_props.append(cell_prop)
            cells_props.append(row_props)
        return cells_props

    def get_all_properties(self):
        return {
            'alignment': self.get_alignment(),
            'allow_autofit': self.get_allow_autofit(),
            'table_width': self.get_table_width(),
            'borders': self.get_table_borders(),
            'table_look': self.get_table_look(),
            'columns_width': self.get_columns_width(),
            'rows_height': self.get_rows_height(),
            'cells_properties': self.get_cells_properties()
        }

# Hàm để lấy thông tin tất cả các bảng trong tài liệu
def extract_tables_properties(docx_file):
    doc = Document(docx_file)
    tables_info = []
    for table in doc.tables:
        helper = TableHelper(table)
        table_info = helper.get_all_properties()
        tables_info.append(table_info)
    return tables_info

# Hàm để chuyển đổi bảng thành HTML với xử lý đầy đủ các thuộc tính
def convert_word_table_to_html(docx_file):
    doc = Document(docx_file)
    html = ['<html><body>']
    
    for table in doc.tables:
        helper = TableHelper(table)
        table_props = helper.get_all_properties()
        
        # Table properties
        alignment = table_props['alignment']
        allow_autofit = table_props['allow_autofit']
        table_width_info = table_props['table_width']
        table_borders = table_props['borders']
        table_look = table_props['table_look']
        columns_width = table_props['columns_width']
        rows_height = table_props['rows_height']
        cells_properties = table_props['cells_properties']
        
        # Tính tổng chiều rộng của tất cả các cột
        table_width = table_width_info['width']
        table_width_pixels = twips_to_pixels(table_width)
        
        # Xử lý căn chỉnh bảng
        if alignment == 'center':
            table_style = f'width: {table_width_pixels}px; margin-left: auto; margin-right: auto; border-collapse: collapse;'
        elif alignment == 'right':
            table_style = f'width: {table_width_pixels}px; margin-left: auto; border-collapse: collapse;'
        else:  # Left or None
            table_style = f'width: {table_width_pixels}px; border-collapse: collapse;'
        
        # Xử lý borders nếu có
        if table_borders:
            border_styles = []
            for side, props in table_borders.items():
                border_style = f'{side}-border: {props["sz"]}pt {props["val"]} #{props["color"] if props["color"] else "000000"};'
                border_styles.append(border_style)
            border_style_str = ' '.join(border_styles)
            table_style += border_style_str
        
        html.append(f'<table style="{table_style}">')
        
        # Xử lý hàng và ô
        for row_idx, row in enumerate(table.rows):
            row_info = rows_height[row_idx]
            if row_info['height']:
                row_height_pixels = row_info['height']
                row_height_rule = row_info['height_rule']
                tr_style = f'height: {row_height_pixels}px;'
            else:
                tr_style = ''
            html.append(f'<tr style="{tr_style}">')
            
            for col_idx, cell in enumerate(row.cells):
                cell_info = cells_properties[row_idx][col_idx]
                cell_width = cell_info.get('width')
                vertical_alignment = cell_info.get('vertical_alignment')
                grid_span = cell_info.get('grid_span', 1)
                v_merge = cell_info.get('v_merge')
                
                # Xử lý rowspan và colspan
                rowspan = 1
                colspan = grid_span if grid_span > 1 else 1
                
                # Kiểm tra v_merge để xử lý rowspan
                if v_merge == 'restart':
                    # Tìm số hàng mà ô này span
                    span_count = 1
                    for r in range(row_idx + 1, len(cells_properties)):
                        if cells_properties[r][col_idx].get('v_merge') == 'continue':
                            span_count += 1
                        else:
                            break
                    rowspan = span_count
                
                # Xử lý chiều rộng ô
                if cell_width:
                    width_style = f'width: {cell_width}px;'
                else:
                    width_style = 'width: auto;'
                
                # Xử lý căn chỉnh dọc
                if vertical_alignment:
                    if vertical_alignment.lower() == 'center':
                        valign_style = 'vertical-align: middle;'
                    elif vertical_alignment.lower() == 'bottom':
                        valign_style = 'vertical-align: bottom;'
                    elif vertical_alignment.lower() == 'top':
                        valign_style = 'vertical-align: top;'
                    else:
                        valign_style = ''
                else:
                    valign_style = ''
                
                # Xử lý borders ô nếu có
                cell_borders = cell_info.get('borders', {})
                if cell_borders:
                    cell_border_styles = []
                    for side, props in cell_borders.items():
                        border_style = f'{side}-border: {props["sz"]}pt {props["val"]} #{props["color"] if props["color"] else "000000"};'
                        cell_border_styles.append(border_style)
                    cell_border_style_str = ' '.join(cell_border_styles)
                else:
                    cell_border_style_str = 'border: 1px solid black;'  # Mặc định
                
                # Xử lý colspan và rowspan
                colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
                rowspan_attr = f' rowspan="{rowspan}"' if rowspan > 1 else ''
                
                # Xử lý nội dung ô
                cell_content = ''
                for paragraph in cell.paragraphs:
                    para_content = ''
                    for run in paragraph.runs:
                        style = ''
                        if run.bold:
                            style += 'font-weight: bold;'
                        if run.italic:
                            style += 'font-style: italic;'
                        if run.font.color.rgb:
                            style += f'color: #{run.font.color.rgb};'
                        para_content += f'<span style="{style}">{run.text}</span>'
                    cell_content += f'<p style="margin: 0;">{para_content}</p>'
                
                # Xây dựng thẻ td
                td_style = f'{width_style} {valign_style} {cell_border_style_str}'
                html.append(f'<td style="{td_style}"{rowspan_attr}{colspan_attr}>{cell_content}</td>')
            html.append('</tr>')
        html.append('</table><br/>')  # Thêm <br/> để tách các bảng nếu có nhiều bảng
    html.append('</body></html>')
    return ''.join(html)

# Ví dụ sử dụng
if __name__ == "__main__":
    docx_path = 'input.docx'
    tables_properties = extract_tables_properties(docx_path)
    
    # Xuất ra file JSON để dễ dàng kiểm tra
    with open('tables_properties.json', 'w', encoding='utf-8') as f:
        json.dump(tables_properties, f, ensure_ascii=False, indent=4)
    
    # Chuyển đổi bảng thành HTML
    html_output = convert_word_table_to_html(docx_path)
    with open('output.html', 'w', encoding='utf-8') as f:
        f.write(html_output)
    
    # In ra màn hình (tuỳ chọn)
    print(json.dumps(tables_properties, ensure_ascii=False, indent=4))
