# services/docx_converter.py

import uuid
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import base64
from docx.shared import RGBColor
from bs4 import BeautifulSoup
import re
from docx.image.exceptions import UnrecognizedImageError
import logging
from app.helpers.table_helper import TableHelper
from lxml import etree
from app.helpers.table_converter_helper import apply_table_styles, apply_cell_styles, twips_to_pixels

logger = logging.getLogger(__name__)

def twips_to_pixels(twips):
    """
    Chuyển đổi từ đơn vị twips (dxa) sang pixel.
    1 twip = 1/20 điểm, và 1 điểm = 1.333 pixel (cho 96 DPI).
    Công thức:
    pixels = twips * (1 / 20) * 1.333
    """
    return twips * (1 / 20) * 1.333

def is_header_row(row):
    """
    Kiểm tra <w:tblHeader> trong XML của row.
    Nếu row có <w:tblHeader>, ta coi row này là header.
    """
    tr_element = row._element
    trPr = tr_element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr')
    if trPr is not None:
        tblHeader = trPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader')
        if tblHeader is not None:
            return True
    return False

class DocxToHtmlConverter:
    def __init__(self, input_path):
        self.doc = Document(input_path)
        self.soup = BeautifulSoup('<html><head></head><body></body></html>', 'html.parser')
        self.list_stack = []
        self.nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    def escape_html(self, text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')

    def get_style_properties(self, run):
        try:
            properties = {}
            font = run.font
            
            properties['font_size'] = font.size.pt if font.size else None
            properties['bold'] = font.bold
            properties['italic'] = font.italic
            properties['underline'] = font.underline
            properties['strikethrough'] = font.strike
            properties['font_name'] = font.name
            # properties['font_family'] = font.family
            properties['small_caps'] = font.small_caps
            properties['all_caps'] = font.all_caps
            properties['superscript'] = font.superscript
            properties['subscript'] = font.subscript

            if font.highlight_color:
                properties['highlight_color'] = font.highlight_color
            else:
                properties['highlight_color'] = None

            if font.color and font.color.rgb:
                try:
                    color = font.color.rgb
                    if isinstance(color, RGBColor) and hasattr(color, 'r') and hasattr(color, 'g') and hasattr(color, 'b'):
                        rgb_int = (color.r << 16) + (color.g << 8) + color.b
                        properties['color'] = rgb_int
                    else:
                        properties['color'] = None
                except AttributeError:
                    properties['color'] = None
            else:
                properties['color'] = None
            return properties
        except Exception as e:
            logger.error(f"An error occurred get_style_properties: {e}")
            return None

    def build_style_str(self, style_dict):
        """
        Từ dictionary mô tả style (font_size, bold, italic,...),
        trả về một chuỗi style hợp lệ (hoặc None nếu không có thuộc tính).
        """
        style_parts = []

        if style_dict.get('font_size'):
            style_parts.append(f'font-size: {style_dict["font_size"]}pt')
        if style_dict.get('bold'):
            style_parts.append('font-weight: bold')
        if style_dict.get('italic'):
            style_parts.append('font-style: italic')
        if style_dict.get('underline'):
            style_parts.append('text-decoration: underline')
        if style_dict.get('color'):
            # color là kiểu int (RGB), format thành hex 6 ký tự
            style_parts.append(f'color: #{style_dict["color"]:06X}')

        # Nếu không có thuộc tính style nào, trả về None
        if style_parts:
            return "; ".join(style_parts)
        return None

    def group_runs_by_style(self, paragraph):
        try:
            runs = paragraph.runs
            groups = []
            if not runs:
                return groups
            
            def clean_text(text):
                """Hàm làm sạch văn bản"""
                try:
                    return (
                        text.replace('…', '...')
                        .replace(' ', ' ')  # Thay khoảng trắng không phải ASCII bằng khoảng trắng chuẩn
                        .replace('. .', '...')
                        .replace('.. ', '...')
                        .replace(' ..', '...')
                        .replace('./.', '...')
                        .replace('/ ..', '...')
                        .replace('/..', '...')
                        .replace('./ ', '...')
                        .replace('/ .', '...')
                        .replace(' ..', '...')
                        .replace('\t', '...')
                        # .replace('  ', ' ')  # Loại bỏ khoảng trắng thừa
                        # .strip()  # Xóa khoảng trắng đầu và cuối chuỗi
                    )
                except Exception as e:
                    logger.error(f"An error occurred: {e}")
                    return text

            current_group = {
                'style': self.get_style_properties(runs[0]),
                'text': clean_text(self.escape_html(runs[0].text)),
                # LƯU LẠI RUN ĐỂ DÙNG Ở convert_paragraph
                'run': runs[0]
            }
            
            for run in runs[1:]:
                # Debug tùy ý
                run_style = self.get_style_properties(run)
                run_text = clean_text(self.escape_html(run.text))

                # So sánh style cũ với run_style mới
                # Nếu cùng style và run_text không chứa \n => gộp
                # Ngược lại => tách group
                if (run_style == current_group['style']
                    and '\n' not in run_text
                    and '\r' not in run_text
                    ):
                    current_group['text'] += run_text
                    current_group['text'] = clean_text(current_group['text'])
                else:
                    groups.append(current_group)
                    current_group = {
                        'style': run_style,
                        'text': run_text,
                        'run': run
                    }

            # Append group cuối
            groups.append(current_group)

            return groups
        except Exception as e:
            logger.error(f"An error occurred group_runs_by_style: {e}")
            return None

    def handle_hyperlinks(self, paragraph, groups):
        try:
            hyperlinks = paragraph._element.findall('.//w:hyperlink', namespaces=self.nsmap)
            for hyperlink in hyperlinks:
                runs = hyperlink.findall('.//w:r', namespaces=self.nsmap)
                text = ''.join([self.escape_html(run.text) for run in runs])
                href = hyperlink.get(qn('r:href'))
                for group in groups:
                    if group['text'] == text:
                        group['text'] = f'<a href="{href}">{text}</a>'
            return groups
        except Exception as e:
            logger.error(f"An error occurred handle_hyperlinks: {e}")
            return None

    def convert_paragraph(self, paragraph):
        try:
            # 1) Gom nhóm runs có cùng style
            groups = self.group_runs_by_style(paragraph)

            # 2) Xử lý hyperlink
            groups = self.handle_hyperlinks(paragraph, groups)

            # Tạo p_tag với text-align
            p_tag = self.soup.new_tag(
                'p',
                style=f'text-align: {paragraph.alignment.name.lower()};'
                if paragraph.alignment else 'text-align: left;'
            )

            for group in groups:
                if group['text'].strip() == '':
                    # Nếu nhóm có text trống, chèn <br/>
                    br_tag = self.soup.new_tag('br')
                    p_tag.append(br_tag)
                    continue  # Bỏ qua phần xử lý text
                # -- Build style_str (giữ nguyên logic cũ) --
                style_str = []
                if group['style']['font_size']:
                    style_str.append(f'font-size: {group["style"]["font_size"]}pt')
                if group['style']['bold']:
                    style_str.append('font-weight: bold')
                if group['style']['italic']:
                    style_str.append('font-style: italic')
                if group['style']['underline']:
                    style_str.append('text-decoration: underline')
                if group['style']['color']:
                    style_str.append(f'color: #{group["style"]["color"]:06X}')

                final_style = '; '.join(style_str) if style_str else None

                # if group['text'].strip() == '':
                #     group['text'] = "\n"

                # text ở đây có thể chứa \n hoặc \r (xuống dòng mềm)
                text = group['text']

                # Tách text theo dòng (khi SHIFT+ENTER => Word gộp vào .text có ký tự \n)
                # splitlines(keepends=False) -> cắt \n, \r\n ra
                lines = text.splitlines()

                # Duyệt từng dòng, với logic regex(\.{3,}) cũ
                for idx, line_content in enumerate(lines):
                    # --- Áp regex tìm '...' ---
                    pattern = r'(\.{3,})'
                    matches = list(re.finditer(pattern, line_content))

                    last_end = 0

                    # Tạo 1 <span> cho dòng này
                    line_span = self.soup.new_tag('span')
                    if final_style:
                        line_span['style'] = final_style

                    for match in matches:
                        start, end = match.span()
                        if start > last_end:
                            sub_text = line_content[last_end:start]
                            line_span.append(sub_text)

                        dots_sub = self.soup.new_tag('span', id=str(uuid.uuid4()))
                        if final_style:
                            dots_sub['style'] = final_style
                        dots_sub.string = match.group(1)
                        line_span.append(dots_sub)
                        last_end = end

                    # Đoạn còn lại sau match cuối
                    if last_end < len(line_content):
                        remaining_text = line_content[last_end:]
                        line_span.append(remaining_text)

                    # Đưa line_span vào p_tag
                    p_tag.append(line_span)

                    # Nếu chưa phải dòng cuối => chèn <br> để xuống dòng
                    # (vì SHIFT+ENTER => "xuống dòng mềm" trong cùng paragraph)
                    if idx < len(lines) - 1:
                        br_tag = self.soup.new_tag('br')
                        p_tag.append(br_tag)

            return p_tag

        except Exception as e:
            logger.error(f"An error occurred convert_paragraph: {e}")
            return None

    
    def is_list_paragraph(self, paragraph):
        try:
            pPr = paragraph._element.pPr
            if pPr is not None and pPr.numPr is not None and pPr.numPr.numId is not None:
                return True
            return False
        except Exception as e:
            logger.error(f"An error occurred is_list_paragraph: {e}")
            return None

    def convert_list_item(self, paragraph):
        try:
            li_tag = self.soup.new_tag('li')
            li_tag.string = paragraph.text
            return li_tag
        except Exception as e:
            logger.error(f"An error occurred convert_list_item: {e}")
            return None

    def handle_list(self, list_item):
        try:
            if not self.list_stack:
                ul_tag = self.soup.new_tag('ul')
                ul_tag.append(list_item)
                self.soup.body.append(ul_tag)
                self.list_stack.append(ul_tag)
            else:
                self.list_stack[-1].append(list_item)
            return list_item
        except Exception as e:
            logger.error(f"An error occurred handle_list: {e}")
            return None

    def convert_headings(self, paragraph):
        try:
            if paragraph.style and paragraph.style.type == WD_STYLE_TYPE.PARAGRAPH and paragraph.style.name.startswith('Heading'):
                # Remove both "Heading " and "#" from the style name
                heading_level_str = paragraph.style.name.replace('Heading ', '').replace('#', '').strip()
                level = int(heading_level_str)
                heading_tag = f'h{level}'
                h_tag = self.soup.new_tag(heading_tag)
                h_tag.string = paragraph.text
                return h_tag
            return None
        except Exception as e:
            logger.error(f"An error occurred convert_headings: {e}")
            return None

    def convert_tables(self, table):
        """
        Chuyển docx.table.Table -> <table> HTML
        Tích hợp logic chia thead/tbody (is_header_row),
        in ra tblPr, cell attrs, ...
        """
        try:
            # 1) Tạo table_tag
            table_tag = self.soup.new_tag('table')

            # 2) Lấy property bảng bằng TableHelper
            helper = TableHelper(table)
            table_props = helper.get_all_properties()
            alignment = table_props['alignment']
            allow_autofit = table_props['allow_autofit']
            table_width_info = table_props['table_width']
            table_look = table_props['table_look']
            columns_width = table_props['columns_width']
            rows_height = table_props['rows_height']
            cells_properties = table_props['cells_properties']

            # 2.a) Gọi apply_table_styles(table_tag, table._element.tblPr)
            tblPr = table._element.tblPr
            apply_table_styles(table_tag, tblPr)

            # 3) Nếu table_width_type = dxa => set width px
            table_width_type = table_width_info.get('type')
            table_width_value = table_width_info.get('width')
            if table_width_type == 'dxa' and table_width_value:
                table_width_px = twips_to_pixels(table_width_value)
                # Append vào style
                table_tag['style'] = (table_tag.get('style','') + f' width: {table_width_px}px;')
            elif table_width_type == 'pct' and table_width_value:
                # example
                table_tag['style'] = (table_tag.get('style','') + ' width: 100%;')
            # else: default => do apply_table_styles or existing

            # 3.a) set alignment => center, right...
            if alignment == 'center':
                table_tag['style'] += ' margin-left: auto; margin-right: auto;'
            elif alignment in ['right','end']:
                table_tag['style'] += ' margin-left: auto;'

            # 3.b) Xử lý autofit => colgroup
            if not allow_autofit and columns_width and len(columns_width) > 0:
                total_width_px = sum(columns_width)
                if total_width_px > 0:
                    colgroup_tag = self.soup.new_tag('colgroup')
                    for w in columns_width:
                        pct = (w / total_width_px)*100 if w else 0
                        col_tag = self.soup.new_tag('col')
                        col_tag['style'] = f'width:{pct:.2f}%;'
                        colgroup_tag.append(col_tag)
                    table_tag.insert(0, colgroup_tag)

            # 4) Tách thead/tbody
            rows = list(table.rows)
            if not rows:
                return table_tag

            head_rows = []
            body_rows = []
            for row in rows:
                if is_header_row(row):
                    head_rows.append(row)
                else:
                    body_rows.append(row)

            # Tạo thead
            if head_rows:
                thead = self.soup.new_tag('thead')
                for row in head_rows:
                    tr_tag = self._convert_row(row, table, cells_properties, rows_height, is_header=True)
                    thead.append(tr_tag)
                table_tag.append(thead)

            # Tạo tbody
            if body_rows:
                tbody = self.soup.new_tag('tbody')
                for row in body_rows:
                    tr_tag = self._convert_row(row, table, cells_properties, rows_height, is_header=False)
                    tbody.append(tr_tag)
                table_tag.append(tbody)

            # Thêm table-look => class
            if table_look:
                look_val = table_look.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if look_val:
                    table_tag['class'] = f'table-look-{look_val}'

            return table_tag

        except Exception as e:
            logger.error(f"Error in convert_tables: {e}")
            return None



    def _convert_row(self, row, table, cells_properties, rows_height, is_header=False):
        """
        row: docx.table._Row
        is_header: True => <th>, False => <td>
        """
        tr_tag = self.soup.new_tag('tr')
        row_index = row._index

        # set row height
        height_info = rows_height[row_index]
        if height_info and height_info['height']:
            tr_tag['style'] = f"height: {height_info['height']}px;"

        num_cells = len(row.cells)

        for col_idx, cell in enumerate(row.cells):
            cell_info = cells_properties[row_index][col_idx]
            tag_name = 'th' if is_header else 'td'
            td_tag = self.soup.new_tag(tag_name)

            # colspan
            colspan = cell_info.get('grid_span',1)
            if colspan > 1:
                td_tag['colspan'] = str(colspan)

            # v_merge => rowspan
            v_merge = cell_info.get('v_merge')
            rowspan = 1
            if v_merge == 'restart':
                # Tính số row
                span_count = 1
                for r2 in range(row_index+1, len(table.rows)):
                    next_cell_info = cells_properties[r2][col_idx]
                    if next_cell_info.get('v_merge') == 'continue':
                        span_count += 1
                    else:
                        break
                if span_count > 1:
                    td_tag['rowspan'] = str(span_count)
            elif v_merge == 'continue':
                # Ẩn ô => skip hiển thị => or style="display:none"
                # Ở đây ta skip hẳn
                continue

            # Gán cell style ban đầu => width: X%
            # (có code chia cột, v.v.)
            # Hoặc let it auto
            # => Ta tạm bỏ qua logic percentage row

            # Gọi apply_cell_styles
            apply_cell_styles(td_tag, cell)

            # Xử lý nội dung paragraphs
            for paragraph in cell.paragraphs:
                p_html = self.convert_paragraph(paragraph)
                self.append_elements(td_tag, p_html)

            # Xử lý nested table
            for nested_table in cell.tables:
                nested_html = self.convert_tables(nested_table)
                self.append_elements(td_tag, nested_html)

            tr_tag.append(td_tag)

        return tr_tag

        
    def append_elements(self, parent, elements):
        try:
            if elements is None:
                return  # Không chèn gì nếu elements là None

            if isinstance(elements, list):
                for el in elements:
                    if el is not None:
                        parent.append(el)
            else:
                parent.append(elements)
        except Exception as e:
            logger.error(f"An error occurred append_elements: {e}")
            return None

    def convert_document(self):
        try:
            html = BeautifulSoup('', 'html.parser')
            head = html.new_tag('head')

            meta_tag = html.new_tag('meta', charset='UTF-8')
            head.append(meta_tag)

            style = html.new_tag('style')
            style.string = """
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1, h2, h3, h4, h5, h6 { color: #2e6c80; }
                table, th, td { padding: 20px; text-align: left; }
                img { max-width: 100%; height: auto; }
                ul, ol { margin: 0; padding-left: 40px; }
                td {vertical-align: top;}
                p { margin: 0 0 1em 0; }
            """
            head.append(style)
            html.append(head)
            body = html.new_tag('body')
            html.append(body)

            if self.doc.sections:
                first_section = self.doc.sections[0]
                top_pt = first_section.top_margin.pt
                right_pt = first_section.right_margin.pt
                bottom_pt = first_section.bottom_margin.pt
                left_pt = first_section.left_margin.pt

                # Chuyển points -> px (1 pt ~ 1.3333 px)
                top_px = int(top_pt * 1.3333)
                right_px = int(right_pt * 1.3333)
                bottom_px = int(bottom_pt * 1.3333)
                left_px = int(left_pt * 1.3333)

                # Gán margin vào body style
                margin_style = f"margin: {top_px}px {right_px}px {bottom_px}px {left_px}px;"
                body['style'] = margin_style

            for element in self.doc._element.body.iterchildren():
                if element.tag == f'{{{self.nsmap["w"]}}}p':
                    paragraph = Paragraph(element, self.doc)
                    if self.is_list_paragraph(paragraph):
                        list_item = self.convert_list_item(paragraph)
                        self.handle_list(list_item)
                    else:
                        heading = self.convert_headings(paragraph)
                        if heading:
                            self.append_elements(body, heading)
                        else:
                            p_tags = self.convert_paragraph(paragraph)
                            self.append_elements(body, p_tags)
                elif element.tag == f'{{{self.nsmap["w"]}}}tbl':
                    table = Table(element, self.doc)
                    table_tag = self.convert_tables(table)
                    self.append_elements(body, table_tag)
                elif element.tag == f'{{{self.nsmap["w"]}}}drawing':
                    # Xử lý drawing => image
                    for inline in element.findall('.//w:inline', namespaces=self.nsmap):
                        try:
                            image_part_id = inline.find('.//wp:docPr', namespaces=self.nsmap).get('id')
                            image_part = self.doc.part.related_parts[image_part_id]
                            blob = image_part.blob
                            b64 = base64.b64encode(blob).decode('utf-8')
                            img_tag = self.soup.new_tag('img', src=f"data:{image_part.content_type};base64,{b64}")
                            self.append_elements(body, img_tag)
                        except UnrecognizedImageError:
                            pass

            return html.prettify()
        except Exception as e:
            logger.error(f"An error occurred convert_document: {e}")
            return None
