import re
import io
import time
import logging
from typing import List
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed

from .html_to_json_service import HtmlToJsonService

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
AiModel = HtmlToJsonService()

class ProcessFileService:
    def __init__(self):
        self.MAX_RETRIES = 3
        self.RETRY_DELAY = 5  # giây
        self.MAX_WORKERS = 5  # Số luồng tối đa

    def _create_prompt(self, check_type: str, text: str, max_length: int = 2000) -> str:
        prompts = {
            "spelling_grammar": (
                "Hãy kiểm tra lỗi chính tả và ngữ pháp trong văn bản sau đây bằng tiếng Việt. \n"
                "Đây là văn bản công việc của cơ quan, yêu cầu văn phong nghiêm túc và chuyên nghiệp.\n"
                "Tập trung vào:\n"
                "1. Lỗi chính tả: Chỉ ra và sửa các lỗi chính tả.\n"
                "2. Lỗi ngữ pháp: Chỉ ra và sửa các lỗi về cú pháp, thì của động từ, và sự phù hợp giữa chủ ngữ và vị ngữ.\n"
                "Đưa ra nhận xét cho từng lỗi, bắt đầu mỗi nhận xét bằng '[NHẬN XÉT]:' và kết thúc bằng dấu '.'.\n"
                "Trước mỗi nhận xét, hãy trích dẫn đoạn văn bản có lỗi, bắt đầu bằng '[TRÍCH DẪN]:' và kết thúc bằng dấu '.'"
            ),
            "content_suggestion": (
                "Hãy đưa ra gợi ý chỉnh sửa nội dung cho văn bản sau đây bằng tiếng Việt. \n"
                "Đây là văn bản công việc của cơ quan, yêu cầu văn phong nghiêm túc và chuyên nghiệp.\n"
                "Tập trung vào:\n"
                "1. Cấu trúc văn bản: Đánh giá và đề xuất cải thiện về bố cục, sự liên kết giữa các phần.\n"
                "2. Tính logic và mạch lạc: Nhận xét về sự nhất quán và hợp lý của các ý tưởng, đề xuất cách cải thiện.\n"
                "3. Độ chính xác và đầy đủ của thông tin: Chỉ ra những phần cần bổ sung hoặc làm rõ thêm.\n"
                "4. Phong cách viết: Đánh giá và gợi ý cải thiện về ngôn ngữ, giọng điệu phù hợp với mục đích văn bản.\n"
                "5. Tính thuyết phục: Đề xuất cách tăng cường sức mạnh thuyết phục của văn bản (nếu cần).\n"
                "Đưa ra nhận xét cho từng vấn đề, bắt đầu mỗi nhận xét bằng '[NHẬN XÉT]:' và kết thúc bằng dấu '.'.\n"
                "Trước mỗi nhận xét, hãy trích dẫn đoạn văn bản liên quan, bắt đầu bằng '[TRÍCH DẪN]:' và kết thúc bằng dấu '.'"
            )
        }

        # Giới hạn độ dài của văn bản đầu vào
        if len(text) > max_length:
            text = text[:max_length] + "..."

        prompt = prompts[check_type] + f"\n\nVăn bản cần kiểm tra:\n\n{text}"
        return prompt

    def _clean_comment(self, comment: str) -> str:
        cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', comment)
        return cleaned.strip()

    def _extract_comments(self, text: str) -> List[str]:
        comments = []
        pattern = r'\[TRÍCH DẪN\]:(.*?)\[NHẬN XÉT\]:(.*?)(?=\[TRÍCH DẪN\]|\Z)'
        matches = re.findall(pattern, text, re.DOTALL)

        logger.debug(f"Số lượng matches tìm thấy: {len(matches)}")

        for match in matches:
            citation = self._clean_comment(match[0])
            comment = self._clean_comment(match[1])
            comments.append(f"[TRÍCH DẪN]: {citation}\n[NHẬN XÉT]: {comment}")

        if not comments:
            logger.warning(f"Không tìm thấy nhận xét nào. Nội dung phản hồi (phần đầu):\n{text[:500]}...")
        return comments

    def _process_prompt(self, check_type: str, full_text: str, selected_model: str) -> List[str]:
        for attempt in range(self.MAX_RETRIES):
            try:
                logger.info(f"Đang xử lý {check_type}...")
                prompt = self._create_prompt(check_type, full_text)
                logger.info(f"Đã tạo prompt cho {check_type}. Độ dài: {len(prompt)} ký tự")

                logger.info(f"Đang gửi yêu cầu kiểm tra {check_type} đến API...")
                response = AiModel.generate_content(prompt, selected_model)

                if response is None or response.startswith("ERROR:"):
                    if response and "429 Resource has been exhausted" in response:
                        logger.info(f"Quota API đã hết, chờ {self.RETRY_DELAY} giây rồi thử lại...")
                        time.sleep(self.RETRY_DELAY)
                        continue
                    raise Exception(response or f"Không nhận được phản hồi từ API cho {check_type}")

                logger.info(f"Đã nhận phản hồi từ API cho {check_type}. Độ dài: {len(response)} ký tự")
                comments = self._extract_comments(response)
                logger.info(f"Đã tạo {len(comments)} nhận xét cho {check_type}")
                return comments
            except Exception as e:
                logger.error(f"Lỗi khi xử lý {check_type}: {str(e)}")
                if attempt < self.MAX_RETRIES - 1:
                    logger.info(f"Thử lại lần {attempt + 1}/{self.MAX_RETRIES}...")
                else:
                    logger.error("Đã thử quá số lần quy định.")
        return []

    def _process_single_file(
        self,
        file_bytes: bytes,
        spelling_grammar: bool,
        content_suggestion: bool,
        selected_model: str,
    ) -> str:
        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            logger.info(f"Đã đọc nội dung file. Độ dài văn bản: {len(full_text)} ký tự")
        except Exception as e:
            raise Exception(f"Lỗi khi đọc file: {str(e)}")

        all_comments = []
        check_types = []
        if spelling_grammar:
            check_types.append("spelling_grammar")
        if content_suggestion:
            check_types.append("content_suggestion")

        with ThreadPoolExecutor(max_workers=len(check_types) or 1) as executor:
            future_to_check = {
                executor.submit(self._process_prompt, check_type, full_text, selected_model): check_type
                for check_type in check_types
            }
            for future in as_completed(future_to_check):
                check_type = future_to_check[future]
                try:
                    comments = future.result()
                    if comments:
                        if check_type == "spelling_grammar":
                            all_comments.append("Spelling and Grammar")
                            all_comments.append("-" * 40)
                            all_comments.extend(comments)
                            all_comments.append("\n")
                        elif check_type == "content_suggestion":
                            all_comments.append("Content Suggestions")
                            all_comments.append("-" * 40)
                            all_comments.extend(comments)
                            all_comments.append("\n")
                    else:
                        logger.info(f"Không có nhận xét nào được tạo cho {check_type}")
                except Exception as exc:
                    logger.error(f"Lỗi khi xử lý {check_type}: {str(exc)}")

        if not all_comments:
            logger.warning("Không tìm thấy nhận xét nào trong phản hồi")
            return "Không có nhận xét nào được tạo ra cho file này."
        else:
            return "\n".join(all_comments)

    def process_question_logic(
        self,
        file_bytes: bytes,
        question: str,
        selected_model: str,
    ) -> str:
        try:
            doc = Document(io.BytesIO(file_bytes))
            content = "\n".join([para.text for para in doc.paragraphs])
            prompt = (
                f"Dưới đây là nội dung của một văn bản:\n\n{content}\n\n"
                f"Câu hỏi: {question}\n\n"
                "Hãy trả lời câu hỏi trên dựa trên nội dung văn bản. "
                "Nếu câu hỏi không liên quan đến nội dung văn bản, hãy trả lời rằng câu hỏi không liên quan."
            )
            response = AiModel.generate_content(prompt, selected_model)
            if response is None or response.startswith("ERROR:"):
                raise Exception(response or "Không nhận được phản hồi từ API.")
            return self._clean_comment(response)
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý câu hỏi: {str(e)}")
